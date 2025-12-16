import os
import io
import json
import shutil
import tempfile
import time
import hashlib
import xml.etree.ElementTree as ET
from typing import List, Tuple

import streamlit as st
import pandas as pd

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from openai import RateLimitError, APIError, APITimeoutError

from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter


# =========================
# App Config + Styling
# =========================
st.set_page_config(page_title="BizChat", layout="wide")

st.markdown(
    """
    <style>
      header {visibility: hidden;}
      footer {visibility: hidden;}
      #MainMenu {visibility: hidden;}
      .block-container {padding-top: 1.2rem; padding-bottom: 1.5rem;}
      .stButton>button {border-radius: 12px; padding: 0.6rem 1rem;}
      section[data-testid="stSidebar"] .block-container {padding-top: 1rem;}
      .bizcard {
        padding: 14px 16px; border-radius: 16px;
        border: 1px solid rgba(49,51,63,0.2);
        margin-bottom: 10px;
      }
    </style>
    """,
    unsafe_allow_html=True,
)

SUPPORTED_EXTS = {
    "pdf", "docx", "xlsx", "xls", "csv", "txt", "json", "xml", "png", "jpg", "jpeg", "webp"
}


# =========================
# Auth (optional)
# =========================
def password_gate():
    """
    If APP_PASSWORD is set in Secrets, require a password.
    """
    app_pwd = st.secrets.get("APP_PASSWORD", "").strip()
    if not app_pwd:
        return

    if "auth_ok" not in st.session_state:
        st.session_state.auth_ok = False

    if st.session_state.auth_ok:
        return

    st.markdown("### 🔒 Team Access")
    pwd = st.text_input("Enter password", type="password")
    if st.button("Login", use_container_width=True):
        if pwd == app_pwd:
            st.session_state.auth_ok = True
            st.rerun()
        else:
            st.error("Wrong password")
    st.stop()


# =========================
# OpenAI Key (BYOK)
# =========================
def get_active_openai_key() -> str:
    """
    BYOK default:
    - User enters key in sidebar (session only)
    - Else fallback to OPENAI_API_KEY in Secrets/env
    """
    if "user_openai_key" not in st.session_state:
        st.session_state.user_openai_key = ""

    with st.sidebar:
        st.subheader("🔐 API Access")

        mode = st.radio(
            "API key mode",
            ["Use my own key (recommended)", "Use company key"],
            index=0,
        )

        if mode == "Use my own key (recommended)":
            st.session_state.user_openai_key = st.text_input(
                "Paste your OpenAI API key",
                type="password",
                placeholder="sk-...",
            ).strip()

            st.caption("Your key is used only in this browser session (not saved).")

            if not st.session_state.user_openai_key:
                st.info("Paste your OpenAI key to continue.")
                st.stop()

            return st.session_state.user_openai_key

        # Company key fallback
        key = st.secrets.get("OPENAI_API_KEY", None) or os.environ.get("OPENAI_API_KEY")
        if not key:
            st.error("Company key not found. Add OPENAI_API_KEY in Streamlit Secrets.")
            st.stop()
        return key


# =========================
# File Parsing
# =========================
def read_pdf_bytes(file_bytes: bytes) -> str:
    parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            t = page.get_text("text")
            if t:
                parts.append(t)
    return "\n".join(parts).strip()

def read_docx_bytes(file_bytes: bytes) -> str:
    bio = io.BytesIO(file_bytes)
    doc = DocxDocument(bio)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paras).strip()

def read_txt_bytes(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore").strip()
    except Exception:
        return file_bytes.decode(errors="ignore").strip()

def read_csv_bytes(file_bytes: bytes) -> str:
    df = pd.read_csv(io.BytesIO(file_bytes))
    return df.to_string(index=False)

def read_excel_bytes(file_bytes: bytes) -> str:
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    out = []
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet)
        out.append(f"--- Sheet: {sheet} ---\n{df.to_string(index=False)}\n")
    return "\n".join(out).strip()

def read_json_bytes(file_bytes: bytes) -> str:
    obj = json.loads(file_bytes.decode("utf-8", errors="ignore"))
    return json.dumps(obj, indent=2, ensure_ascii=False)

def read_xml_bytes(file_bytes: bytes) -> str:
    root = ET.fromstring(file_bytes.decode("utf-8", errors="ignore"))
    lines = []

    def walk(node, depth=0):
        indent = "  " * depth
        tag = node.tag
        text = (node.text or "").strip()
        if text:
            lines.append(f"{indent}<{tag}> {text}")
        else:
            lines.append(f"{indent}<{tag}>")
        for child in list(node):
            walk(child, depth + 1)

    walk(root, 0)
    return "\n".join(lines).strip()

def read_image_ocr_bytes(file_bytes: bytes) -> str:
    img = Image.open(io.BytesIO(file_bytes))
    txt = pytesseract.image_to_string(img)
    return (txt or "").strip()

def file_to_text(filename: str, file_bytes: bytes) -> Tuple[str, str]:
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        return read_pdf_bytes(file_bytes), "pdf"
    if ext == "docx":
        return read_docx_bytes(file_bytes), "docx"
    if ext in ("xlsx", "xls"):
        return read_excel_bytes(file_bytes), "excel"
    if ext == "csv":
        return read_csv_bytes(file_bytes), "csv"
    if ext == "txt":
        return read_txt_bytes(file_bytes), "txt"
    if ext == "json":
        return read_json_bytes(file_bytes), "json"
    if ext == "xml":
        return read_xml_bytes(file_bytes), "xml"
    if ext in ("png", "jpg", "jpeg", "webp"):
        return read_image_ocr_bytes(file_bytes), "image_ocr"
    return "", "unknown"


# =========================
# Indexing + RAG
# =========================
def chunk_documents(docs: List[LCDocument]) -> List[LCDocument]:
    # Larger chunk => fewer embedding calls => fewer rate-limit issues
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1600,
        chunk_overlap=200,
        separators=["\n\n", "\n", " ", ""],
    )
    return splitter.split_documents(docs)

def build_faiss_with_retries(chunks: List[LCDocument], embeddings: OpenAIEmbeddings, max_retries: int = 6):
    wait = 2
    for attempt in range(1, max_retries + 1):
        try:
            return FAISS.from_documents(chunks, embeddings)
        except RateLimitError:
            if attempt == max_retries:
                raise
            st.warning(f"Rate limit hit while embedding. Retrying in {wait}s (attempt {attempt}/{max_retries})...")
            time.sleep(wait)
            wait = min(wait * 2, 30)
        except (APITimeoutError, APIError):
            if attempt == max_retries:
                raise
            st.warning(f"Temporary API error. Retrying in {wait}s (attempt {attempt}/{max_retries})...")
            time.sleep(wait)
            wait = min(wait * 2, 30)

def format_sources(source_docs: List[LCDocument]) -> str:
    out = []
    for i, d in enumerate(source_docs[:6], start=1):
        src = d.metadata.get("source", "unknown")
        snippet = (d.page_content[:250] or "").replace("\n", " ").strip()
        out.append(f"[{i}] {src} — {snippet}...")
    return "\n".join(out).strip()

def build_pdf_bytes(title: str, body: str) -> bytes:
    buff = io.BytesIO()
    c = canvas.Canvas(buff, pagesize=A4)
    width, height = A4

    x = 40
    y = height - 60

    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, title[:120])
    y -= 30

    c.setFont("Helvetica", 10)
    max_chars = 100

    for para in body.split("\n"):
        while len(para) > max_chars:
            line = para[:max_chars]
            para = para[max_chars:]
            if y < 60:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = height - 60
            c.drawString(x, y, line)
            y -= 14

        if y < 60:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 60
        c.drawString(x, y, para)
        y -= 18

    c.save()
    buff.seek(0)
    return buff.read()

def ensure_session_dirs():
    if "temp_dir" not in st.session_state:
        st.session_state.temp_dir = tempfile.mkdtemp(prefix="bizchat_")

def reset_all():
    if "raw_docs" in st.session_state:
        st.session_state.raw_docs = []
    if "vectorstore" in st.session_state:
        st.session_state.vectorstore = None
    if "chat" in st.session_state:
        st.session_state.chat = []
    if "temp_dir" in st.session_state:
        try:
            shutil.rmtree(st.session_state.temp_dir, ignore_errors=True)
        except Exception:
            pass
    st.session_state.temp_dir = tempfile.mkdtemp(prefix="bizchat_")


# =========================
# Page Header
# =========================
password_gate()

col1, col2 = st.columns([0.14, 0.86])
with col1:
    # If you have a local logo in repo, use st.image("logo.png", width=72)
    st.markdown(
        """
        <div class="bizcard">
          <div style="font-weight:700; font-size:16px;">BizChat</div>
          <div style="font-size:12px; opacity:0.8;">RAG Assistant</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
with col2:
    st.markdown("## BizChat — Document Q&A Bot")
    st.caption("Upload files → Build index → Ask questions → Export answers (PDF).")

# =========================
# Sidebar Controls
# =========================
api_key = get_active_openai_key()
os.environ["OPENAI_API_KEY"] = api_key

with st.sidebar:
    st.divider()
    st.subheader("⚙️ Settings")

    model = st.selectbox("Chat model", ["gpt-4o-mini", "gpt-4.1-mini", "gpt-4o"], index=0)
    temperature = st.slider("Temperature", 0.0, 1.0, 0.2, 0.05)

    max_file_mb = st.number_input("Max file size (MB)", min_value=2, max_value=50, value=15, step=1)
    max_chunks = st.number_input("Max chunks to embed", min_value=50, max_value=5000, value=800, step=50)

    st.info(
        "If you get RateLimit:\n"
        "- Use your own API key\n"
        "- Upload fewer/smaller files\n"
        "- Reduce max chunks\n"
        "- Try again after 1–2 minutes"
    )

    st.divider()
    if st.button("🧹 Clear everything", use_container_width=True):
        reset_all()
        st.success("Cleared. Reloading…")
        st.rerun()


# =========================
# Session State
# =========================
ensure_session_dirs()
if "raw_docs" not in st.session_state:
    st.session_state.raw_docs = []
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "chat" not in st.session_state:
    st.session_state.chat = []


# =========================
# Wizard Layout
# =========================
st.markdown("### Step 1 — Add Knowledge")

left, right = st.columns([1.1, 0.9])

with left:
    uploaded = st.file_uploader(
        "Upload PDFs, Word, Excel, CSV, JSON, XML, TXT, Images",
        type=list(SUPPORTED_EXTS),
        accept_multiple_files=True,
    )
    pasted_text = st.text_area("Or paste text here", height=140, placeholder="Paste meeting notes, SOP, emails, etc.")

    add_clicked = st.button("📥 Add to Knowledge Base", use_container_width=True)

    if add_clicked:
        added, failed = 0, 0

        if pasted_text and pasted_text.strip():
            st.session_state.raw_docs.append(
                LCDocument(page_content=pasted_text.strip(), metadata={"source": "pasted_text"})
            )
            added += 1

        if uploaded:
            for f in uploaded:
                try:
                    data = f.read()
                    if len(data) > int(max_file_mb) * 1024 * 1024:
                        st.warning(f"Skipped {f.name}: too large (> {max_file_mb}MB).")
                        failed += 1
                        continue

                    text, src_type = file_to_text(f.name, data)
                    if text and text.strip():
                        out_path = os.path.join(st.session_state.temp_dir, f.name)
                        with open(out_path, "wb") as w:
                            w.write(data)

                        st.session_state.raw_docs.append(
                            LCDocument(page_content=text, metadata={"source": f.name, "type": src_type})
                        )
                        added += 1
                    else:
                        st.warning(f"Skipped {f.name}: no readable text found.")
                        failed += 1
                except Exception:
                    failed += 1

        st.success(f"Added: {added} | Skipped/failed: {failed}")

with right:
    st.markdown("### Step 2 — Build Index")
    st.write(f"Items in memory: **{len(st.session_state.raw_docs)}**")

    build_clicked = st.button(
        "⚡ Build / Rebuild Index (FAISS)",
        use_container_width=True,
        disabled=(len(st.session_state.raw_docs) == 0),
    )

    if build_clicked:
        with st.spinner("Chunking + embedding + building index..."):
            chunks = chunk_documents(st.session_state.raw_docs)

            if len(chunks) > int(max_chunks):
                st.warning(f"Too many chunks ({len(chunks)}). Limiting to {max_chunks} to avoid RateLimit.")
                chunks = chunks[: int(max_chunks)]

            embeddings = OpenAIEmbeddings(model="text-embedding-3-small")

            try:
                st.session_state.vectorstore = build_faiss_with_retries(chunks, embeddings)
                st.success("Index built successfully ✅")
            except RateLimitError:
                st.error(
                    "OpenAI RateLimit while creating embeddings.\n\n"
                    "Fix:\n"
                    "1) Wait 1–2 minutes and click Build again\n"
                    "2) Upload fewer/smaller files\n"
                    "3) Reduce 'Max chunks to embed'\n"
                    "4) Ensure billing/quota is active for this API key"
                )

    if st.session_state.vectorstore is not None:
        st.success("Index status: READY ✅")
    else:
        st.warning("Index status: NOT READY")


st.divider()
st.markdown("### Step 3 — Ask Questions")

if st.session_state.vectorstore is None:
    st.info("Build the index first (Step 2), then ask questions here.")
else:
    # Show chat history
    for msg in st.session_state.chat:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    user_q = st.chat_input("Ask something like: 'Summarize this PDF' or 'What are the key risks?'")
    if user_q:
        st.session_state.chat.append({"role": "user", "content": user_q})
        with st.chat_message("user"):
            st.markdown(user_q)

        llm = ChatOpenAI(model=model, temperature=temperature)
        retriever = st.session_state.vectorstore.as_retriever(search_kwargs={"k": 5})

        with st.spinner("Searching + answering..."):
            source_docs = retriever.get_relevant_documents(user_q)
            context = "\n\n".join(
                [f"Source: {d.metadata.get('source','unknown')}\n{d.page_content}" for d in source_docs]
            )

            prompt = f"""
You are a helpful assistant. Answer the user using ONLY the provided context.
If the answer is not present, say you don't know and suggest what file to upload.

User question:
{user_q}

Context:
{context}

Answer (clear, concise, structured):
"""
            resp = llm.invoke(prompt)
            answer = resp.content.strip()

            sources_text = format_sources(source_docs)
            final_answer = answer
            if sources_text:
                final_answer += "\n\n---\n**Sources (snippets):**\n" + sources_text

        st.session_state.chat.append({"role": "assistant", "content": final_answer})
        with st.chat_message("assistant"):
            st.markdown(final_answer)

        pdf_bytes = build_pdf_bytes("BizChat Answer", final_answer)
        st.download_button(
            "⬇️ Download answer as PDF",
            data=pdf_bytes,
            file_name="bizchat_answer.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
